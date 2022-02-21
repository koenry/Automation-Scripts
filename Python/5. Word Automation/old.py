from datetime import date, timedelta
from docx import Document
from docx import Document

db = "Weekly"
task = "doc2.docx"
doc = Document("doc.docx")
month = int(input('Month?   '))
tw1 = [22, 23, 24, 25, 26, 27, 28]
tw2 = [29, 30, 31]

def monday1(year):
    d = date(year, month, 1)                    
    d += timedelta(days = 7 - d.weekday())
    while d.day < 8:
        yield d
        d += timedelta(days = 7)
        
def monday2(year):
    d = date(year, month, 7)                    
    d += timedelta(days = 7 - d.weekday())
    while  d.day < 15:
        yield d
        d += timedelta(days = 7)

def monday3(year):
    d = date(year, month, 14)                    
    d += timedelta(days = 7 - d.weekday())
    while d.day < 22:
        yield d
        d += timedelta(days = 7)

def monday4(year):
    d = date(year, month, 21)                    
    d += timedelta(days = 7 - d.weekday())
    while d.day in tw1: # < 29 from 2021-06 goes to 2021-08 for some reason? Maybe a bug?
        yield d
        d += timedelta(days = 7)

def monday5(year):
    d = date(year, month, 28)                    
    d += timedelta(days = 7 - d.weekday())
    while d.day in tw2:
        yield d
        d += timedelta(days = 7)



def tuesday1(year):
    d = date(year, month, 1)                    
    d += timedelta(days = 8 - d.weekday())
    while d.day <= 7:
        yield d
        d += timedelta(days = 7)
        
def tuesday2(year):
    d = date(year, month, 7)                    
    d += timedelta(days = 8 - d.weekday())
    while  d.day < 15:
        yield d
        d += timedelta(days = 7)

def tuesday3(year):
    d = date(year, month, 14)                    
    d += timedelta(days = 8 - d.weekday())
    while d.day < 22:
        yield d
        d += timedelta(days = 7)

def tuesday4(year):
    d = date(year, month, 21)                    
    d += timedelta(days = 8 - d.weekday())
    while d.day in tw1: # < 29 from 2021-06 goes to 2021-08 for some reason? Maybe a bug?
        yield d
        d += timedelta(days = 7)

def tuesday5(year):
    d = date(year, month, 28)                    
    d += timedelta(days = 8 - d.weekday())
    while d.day in tw2:
        yield d
        d += timedelta(days = 7)




def wednesday1(year):
    d = date(year, month, 1)                    
    d += timedelta(days = 9 - d.weekday())
    while d.day < 8:
        yield d
        d += timedelta(days = 7)
        
def wednesday2(year):
    d = date(year, month, 7)                    
    d += timedelta(days = 9 - d.weekday())
    while  d.day < 15:
        yield d
        d += timedelta(days = 7)

def wednesday3(year):
    d = date(year, month, 14)                    
    d += timedelta(days = 9 - d.weekday())
    while d.day < 22:
        yield d
        d += timedelta(days = 7)

def wednesday4(year):
    d = date(year, month, 21)                    
    d += timedelta(days = 9 - d.weekday())
    while d.day in tw1: # < 29 from 2021-06 goes to 2021-08 for some reason? Maybe a bug?
        yield d
        d += timedelta(days = 7)

def wednesday5(year):
    d = date(year, month, 28)                    
    d += timedelta(days = 9 - d.weekday())
    while d.day in tw2:
        yield d
        d += timedelta(days = 7)




def thursday1(year):
    d = date(year, month, 1)                    
    d += timedelta(days = 10 - d.weekday())
    while d.day < 8:
        yield d
        d += timedelta(days = 7)
        
def thursday2(year):
    d = date(year, month, 7)                    
    d += timedelta(days = 10 - d.weekday())
    while  d.day < 15:
        yield d
        d += timedelta(days = 7)

def thursday3(year):
    d = date(year, month, 14)                    
    d += timedelta(days = 10 - d.weekday())
    while d.day < 22:
        yield d
        d += timedelta(days = 7)

def thursday4(year):
    d = date(year, month, 21)                    
    d += timedelta(days = 10 - d.weekday())
    while d.day in tw1: # < 29 from 2021-06 goes to 2021-08 for some reason? Maybe a bug?
        yield d
        d += timedelta(days = 7)

def thursday5(year):
    d = date(year, month, 28)                    
    d += timedelta(days = 10 - d.weekday())
    while d.day in tw2:
        yield d
        d += timedelta(days = 7)




def friday1(year):
    d = date(year, month, 1)                    
    d += timedelta(days = 4 - d.weekday())
    while d.day < 8:
        yield d
        d += timedelta(days = 7)
        
def friday2(year):
    d = date(year, month, 7)                    
    d += timedelta(days = 11 - d.weekday())
    while  d.day < 15:
        yield d
        d += timedelta(days = 7)

def friday3(year):
    d = date(year, month, 14)                    
    d += timedelta(days = 11 - d.weekday())
    while d.day < 22:
        yield d
        d += timedelta(days = 7)

def friday4(year):
    d = date(year, month, 21)                    
    d += timedelta(days = 11 - d.weekday())
    while d.day in tw1: # < 29 from 2021-06 goes to 2021-08 for some reason? Maybe a bug?
        yield d
        d += timedelta(days = 7)

def friday5(year):
    d = date(year, month, 28)                    
    d += timedelta(days = 11 - d.weekday())
    while d.day in tw2:
        yield d
        d += timedelta(days = 7)



for d in monday1(2021):
    doc.tables[0].cell(1, 1).text =  str(d)
    doc.tables[0].cell(1, 2).text =  str(task)

for d in monday2(2021):
    doc.tables[0].cell(1, 4).text =  str(d)
    doc.tables[0].cell(1, 5).text =  str(task)

for d in monday3(2021):
    doc.tables[0].cell(1, 7).text =  str(d)
    doc.tables[0].cell(1, 8).text =  str(task)

for d in monday4(2021):
    doc.tables[0].cell(1, 10).text =  str(d)
    doc.tables[0].cell(1, 11).text =  str(task)
    doc.save("grid2.docx")

for d in monday5(2021):
    doc.tables[0].cell(1, 13).text =  str(d) 
    doc.save("grid2.docx")




for d in tuesday1(2021):
    doc.tables[0].cell(1, 16).text =  str(d)
    doc.tables[0].cell(1, 17).text =  str(task)


for d in tuesday2(2021):
    doc.tables[0].cell(1, 19).text =  str(d)
    doc.tables[0].cell(1, 20).text =  str(task)


for d in tuesday3(2021):
    doc.tables[0].cell(1, 22).text =  str(d)
    doc.tables[0].cell(1, 23).text =  str(task)


for d in tuesday4(2021):
    doc.tables[0].cell(1, 25).text =  str(d)
    doc.tables[0].cell(1, 26).text =  str(task)
    doc.save("grid2.docx")            

for d in tuesday5(2021):
    doc.tables[0].cell(1, 28).text =  str(d)
    doc.tables[0].cell(1, 29).text =  str(task)
    doc.save("grid2.docx")


for d in wednesday1(2021):
    doc.tables[0].cell(1, 31).text =  str(d)
    doc.tables[0].cell(1, 32).text =  str(task)
    doc.tables[0].cell(1, 76).text =  str(d)
    doc.tables[0].cell(1, 77).text =  str(db)


for d in wednesday2(2021):
    doc.tables[0].cell(1, 34).text =  str(d)
    doc.tables[0].cell(1, 35).text =  str(task)
    doc.tables[0].cell(1, 79).text =  str(d)
    doc.tables[0].cell(1, 80).text =  str(db)


for d in wednesday3(2021):
    doc.tables[0].cell(1, 37).text =  str(d)
    doc.tables[0].cell(1, 38).text =  str(task)
    doc.tables[0].cell(1, 82).text =  str(d)
    doc.tables[0].cell(1, 83).text =  str(db)


for d in wednesday4(2021):
    doc.tables[0].cell(1, 40).text =  str(d)
    doc.tables[0].cell(1, 41).text =  str(task)
    doc.tables[0].cell(1, 85).text =  str(d)
    doc.tables[0].cell(1, 86).text =  str(db)
    doc.save("grid2.docx")

for d in wednesday5(2021):
    doc.tables[0].cell(1, 43).text =  str(d)
    doc.tables[0].cell(1, 44).text =  str(task)
    doc.tables[0].cell(1, 88).text =  str(d)
    doc.tables[0].cell(1, 89).text =  str(db)
    doc.save("grid2.docx")


for d in thursday1(2021):
    doc.tables[0].cell(1, 46).text =  str(d)
    doc.tables[0].cell(1, 47).text =  str(task)


for d in thursday2(2021):
    doc.tables[0].cell(1, 49).text =  str(d)
    doc.tables[0].cell(1, 50).text =  str(task)


for d in thursday3(2021):
    doc.tables[0].cell(1, 52).text =  str(d)
    doc.tables[0].cell(1, 53).text =  str(task)


for d in thursday4(2021):
    doc.tables[0].cell(1, 55).text =  str(d)
    doc.tables[0].cell(1, 56).text =  str(task)
    doc.save("grid2.docx")

for d in thursday5(2021):
    doc.tables[0].cell(1, 58).text =  str(d)
    doc.tables[0].cell(1, 59).text =  str(task)
    doc.save("grid2.docx")



for d in friday1(2021):
    doc.tables[0].cell(1, 61).text =  str(d)
    doc.tables[0].cell(1, 62).text =  str(task)


for d in friday2(2021):
    doc.tables[0].cell(1, 64).text =  str(d)
    doc.tables[0].cell(1, 65).text =  str(task)


for d in friday3(2021):
    doc.tables[0].cell(1, 67).text =  str(d)
    doc.tables[0].cell(1, 68).text =  str(task)


for d in friday4(2021):
    doc.tables[0].cell(1, 70).text =  str(d)
    doc.tables[0].cell(1, 71).text =  str(task)
    doc.save("grid2.docx")

for d in friday5(2021):
    doc.tables[0].cell(1, 73).text =  str(d)
    doc.tables[0].cell(1, 74).text =  str(task)
    doc.save("grid2.docx")


print("******************************FINISHED******************************")
input("\nPress Enter to continue...")



