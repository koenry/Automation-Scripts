import datetime
from docx import Document


db = "Weekly report."
task = "Daily.docx"
doc = Document("daily.docx")

month = int(input('Which month? (1-12)   '))

monthinteger = month
months = datetime.date(1900, monthinteger, 1).strftime('%B')
year = datetime.datetime.now()

week0 = [1, 2, 3, 4, 5, 6, 7]
week4 = [25, 26, 27,28, 29, 30, 31]

save = 'word file' + months
save2 = save + ".docx"

# monday

def monday0():
    today = datetime.date(year.year, month, 1)
    d = today + datetime.timedelta(days=-today.weekday(), weeks=0)
    if d.day in week0:           # week 0 and week 4 has IN called because the first weekday might overlap with previous month. Ex. may 31st is monday june 1st is tuesday.
        doc.tables[0].cell(1, 1).text =  str(d)
        doc.tables[0].cell(1, 2).text =  str(task)

def monday1():
    today = datetime.date(year.year, month, 1)
    d = today + datetime.timedelta(days=-today.weekday(), weeks=1)
    doc.tables[0].cell(1, 4).text =  str(d)
    doc.tables[0].cell(1, 5).text =  str(task)
    
def monday2():
    today = datetime.date(year.year, month, 2)
    d = today + datetime.timedelta(days=-today.weekday(), weeks=2)
    doc.tables[0].cell(1, 7).text =  str(d)
    doc.tables[0].cell(1, 8).text =  str(task)

def monday3():
    today = datetime.date(year.year, month, 3)
    d = today + datetime.timedelta(days=-today.weekday(), weeks=3)
    doc.tables[0].cell(1, 10).text =  str(d)
    doc.tables[0].cell(1, 11).text =  str(task)



def monday4():
    today = datetime.date(year.year, month, 4)
    d = today + datetime.timedelta(days=-today.weekday(), weeks=4)
    if d.day in week4:
        doc.tables[0].cell(1, 13).text =  str(d)
        doc.tables[0].cell(1, 14).text =  str(task) 




# tuesday



def tuesday0():
    today = datetime.date(year.year, month, 1)
    d = today + datetime.timedelta(days=1 -today.weekday(), weeks=0)
    if d.day in week0:
        doc.tables[0].cell(1, 16).text =  str(d)
        doc.tables[0].cell(1, 17).text =  str(task)


def tuesday1():
    today = datetime.date(year.year, month, 1)
    d = today + datetime.timedelta(days=1 -today.weekday(), weeks=1)
    doc.tables[0].cell(1, 19).text =  str(d)
    doc.tables[0].cell(1, 20).text =  str(task)

def tuesday2():
    today = datetime.date(year.year, month, 2)
    d = today + datetime.timedelta(days=1 -today.weekday(), weeks=2)
    doc.tables[0].cell(1, 22).text =  str(d)
    doc.tables[0].cell(1, 23).text =  str(task)

def tuesday3():
    today = datetime.date(year.year, month, 3)
    d = today + datetime.timedelta(days=1 -today.weekday(), weeks=3)
    doc.tables[0].cell(1, 25).text =  str(d)
    doc.tables[0].cell(1, 26).text =  str(task)
   

def tuesday4():
    today = datetime.date(year.year, month, 4)
    d = today + datetime.timedelta(days=1 -today.weekday(), weeks=4)
    if d.day in week4:
        doc.tables[0].cell(1, 28).text =  str(d)
        doc.tables[0].cell(1, 29).text =  str(task)


# wednesday


def wednesday0():
    today = datetime.date(year.year, month, 1)
    d = today + datetime.timedelta(days=2 -today.weekday(), weeks=0)
    if d.day in week0:
        doc.tables[0].cell(1, 31).text =  str(d)
        doc.tables[0].cell(1, 32).text =  str(task)
        doc.tables[0].cell(1, 76).text =  str(d)
        doc.tables[0].cell(1, 77).text =  str(db)

def wednesday1():
    today = datetime.date(year.year, month, 3)
    d = today + datetime.timedelta(days=2 -today.weekday(), weeks=1)
    doc.tables[0].cell(1, 34).text =  str(d)
    doc.tables[0].cell(1, 35).text =  str(task)
    doc.tables[0].cell(1, 79).text =  str(d)
    doc.tables[0].cell(1, 80).text =  str(db)

def wednesday2():
    today = datetime.date(year.year, month, 3)
    d = today + datetime.timedelta(days=2 -today.weekday(), weeks=2)
    doc.tables[0].cell(1, 37).text =  str(d)
    doc.tables[0].cell(1, 38).text =  str(task)
    doc.tables[0].cell(1, 82).text =  str(d)
    doc.tables[0].cell(1, 83).text =  str(db)

def wednesday3():
    today = datetime.date(year.year, month, 3)
    d = today + datetime.timedelta(days=2 -today.weekday(), weeks=3)
    doc.tables[0].cell(1, 40).text =  str(d)
    doc.tables[0].cell(1, 41).text =  str(task)
    doc.tables[0].cell(1, 85).text =  str(d)
    doc.tables[0].cell(1, 86).text =  str(db)


def wednesday4():
    today = datetime.date(year.year, month, 3)
    d = today + datetime.timedelta(days=2 -today.weekday(), weeks=4)
    if d.day in week4:
        doc.tables[0].cell(1, 43).text =  str(d)
        doc.tables[0].cell(1, 44).text =  str(task)
        doc.tables[0].cell(1, 88).text =  str(d)
        doc.tables[0].cell(1, 89).text =  str(db)



# thursday


def thursday0():
    today = datetime.date(year.year, month, 1)
    d = today + datetime.timedelta(days=3 -today.weekday(), weeks=0)
    if d.day in week0:
        doc.tables[0].cell(1, 46).text =  str(d)
        doc.tables[0].cell(1, 47).text =  str(task)

def thursday1():
    today = datetime.date(year.year, month, 3)
    d = today + datetime.timedelta(days=3 -today.weekday(), weeks=1)
    doc.tables[0].cell(1, 49).text =  str(d)
    doc.tables[0].cell(1, 50).text =  str(task)

def thursday2():
    today = datetime.date(year.year, month, 3)
    d = today + datetime.timedelta(days=3 -today.weekday(), weeks=2)
    doc.tables[0].cell(1, 52).text =  str(d)
    doc.tables[0].cell(1, 53).text =  str(task)

def thursday3():
    today = datetime.date(year.year, month, 3)
    d = today + datetime.timedelta(days=3 -today.weekday(), weeks=3)
    doc.tables[0].cell(1, 55).text =  str(d)
    doc.tables[0].cell(1, 56).text =  str(task)


def thursday4():
    today = datetime.date(year.year, month, 3)
    d = today + datetime.timedelta(days=3 -today.weekday(), weeks=4)
    if d.day in week4:
        doc.tables[0].cell(1, 58).text =  str(d)
        doc.tables[0].cell(1, 59).text =  str(task)

# friday

def friday0():
    today = datetime.date(year.year, month, 1)
    d = today + datetime.timedelta(days=4 -today.weekday(), weeks=0)
    if d.day in week0:
        doc.tables[0].cell(1, 61).text =  str(d)
        doc.tables[0].cell(1, 62).text =  str(task)

def friday1():
    today = datetime.date(year.year, month, 3)
    d = today + datetime.timedelta(days=4 -today.weekday(), weeks=1)
    doc.tables[0].cell(1, 64).text =  str(d)
    doc.tables[0].cell(1, 65).text =  str(task)

def friday2():
    today = datetime.date(year.year, month, 3)
    d = today + datetime.timedelta(days=4 -today.weekday(), weeks=2)
    doc.tables[0].cell(1, 67).text =  str(d)
    doc.tables[0].cell(1, 68).text =  str(task)

def friday3():
    today = datetime.date(year.year, month, 3)
    d = today + datetime.timedelta(days=4 -today.weekday(), weeks=3)
    doc.tables[0].cell(1, 70).text =  str(d)
    doc.tables[0].cell(1, 71).text =  str(task)
    doc.save(save2)


def friday4():
    today = datetime.date(year.year, month, 3)
    d = today + datetime.timedelta(days=4 -today.weekday(), weeks=4)
    if d.day in week4:
        doc.tables[0].cell(1, 73).text =  str(d)
        doc.tables[0].cell(1, 74).text =  str(task)
        doc.save(save2) # double save because if theres no friday on week4 it wont save


monday0()
monday1()
monday2()
monday3()
monday4()

tuesday0()
tuesday1()
tuesday2()
tuesday3()
tuesday4()

wednesday0()
wednesday1()
wednesday2()
wednesday3()
wednesday4()

thursday0()
thursday1()
thursday2()
thursday3()
thursday4()

friday0()
friday1()
friday2()
friday3()
friday4()

print("******************************FINISHED******************************")
input("\nPress Enter to continue...")









